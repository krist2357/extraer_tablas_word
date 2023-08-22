import PySimpleGUI as sg
import docx
from openpyxl import Workbook
from openpyxl import load_workbook
import time
import os

label0 = sg.Text('EXTRACCIÓN DE TABLAS DE WORD Y TRASLADO A EXCEL V1.0')
label1 = sg.Text('LOG DE ACCIONES:')
add_button1 = sg.Button('Listar documentos')
label2 = sg.Text('INGRESE EL NOMBRE DE UN ARCHIVO DE WORD PARA SER LEIDO:')
input_box_1 = sg.InputText(tooltip='Ingrese el archivo de word a leer', key='documento')
add_button2 = sg.Button('Extraer tablas')
logwindow = sg.Output(size=(60,20))
label3 = sg.Text('VERIFICAR EXCEL:')
add_button3 = sg.Button('Verificar Excel')
label4 = sg.Text('INGRESE EL NOMBRE DE LA HOJA DE CÁLCULO:')
input_box_2 = sg.InputText(tooltip='Ingrese el nombre de la hoja de cáclulo', key='hoja')
add_button4 = sg.Button('Trasladar a Excel')
label_5 = sg.Text(' ', key='indication')
label_6 = sg.Text('Desarrollo: @Krist2357, 2023')

window = sg.Window('Traslado de tablas de Word a Excel V1.0', layout=[[label0],
                                                                      [label1],
                                                                      [logwindow], 
                                                                      [add_button1],
                                                                      [label2],
                                                                      [input_box_1, add_button2],
                                                                      [label3, add_button3],
                                                                      [label4],
                                                                      [input_box_2, add_button4],
                                                                      [label_5],
                                                                      [label_6]])

while True:
    event, values = window.read() #Forma un diccionario.
    #print(event)
    #print(values)

    match event:
        case 'Listar documentos':
            dir_actual = os.getcwd() # Listar directorio actual donde está la app.
            new_dir_actual = dir_actual.replace('\\','/') # cambiar el backslash a slash para que sea un directorio 
            list_files = os.listdir(new_dir_actual) #Listar lo que está dentro del directorio.
            print('................................................................')
            print('EL DIRECTORIO DE TRABAJO ES:')
            #print(dir_actual)
            print(new_dir_actual)
            print('................................................................')
            print('LOS ARCHIVOS EXISTENTES EN EL DIRECTORIO:')
            print(list_files)

        case 'Extraer tablas':
            try:
                file = values['documento'] #extraigo del diccionario, el contenido de la llave documento.
                print('................................................................')
                print('EL ARCHIVO A EXTRAER ES:')
                print(file)
                docFile = f'{new_dir_actual}/{file}' # Unir con f-string el archivo que se desea extraer.
                print('................................................................')
                print('LA RUTA DEL ARCHIVO ES:')
                print(docFile)
                #Se usa la libreria docx para tratar los archivos de word:
                doc = docx.Document(docFile)
                # Se filtra las tablas 1, 2, 3, 4, ya que la primera (tabla 0) es la carátula 
                # y las dos últimas son del reporte fotográfico (tablas 5 y 6).
                dict = {i:doc.tables[i] for i in range(len(doc.tables)) if i > 0 and i < 5}
                # Código para extraer los valores de dict, con la llave
                dict2 = {}
                k = 0
                for table_id in dict:
                    #print(table_id, dict[table_id])
                    # extraigo las filas con el método rows
                    for row in dict[table_id].rows:
                        # Guardo en un diccionario la lista generada con los métodos cells y text
                        dict2[k] = [i.text for i in row.cells]
                        k = k + 1
                        #print(dict2)
                print('................................................................')
                print(file)
                print(docFile)
                print(dict)
                print()
                print('!PROCESO DE LECTURA DE TABLAS CORRECTO!')
            except Exception:
                print('Error: No ingresó el nombre correcto del archivo de word.') 
                print('o no es el archivo correcto con extensión .docx')
                window['indication'].update(value = '!!Debe ingresar nombre correcto o archivo correcto¡¡')

        case 'Verificar Excel':
            dir_actual = os.getcwd()
            new_dir_actual = dir_actual.replace('\\','/')
            filesheet = f'{new_dir_actual}/datos.xlsx' #Se requiere tener un archivo de excel creado.
            print('................................................................')
            print('EL ARCHIVO DE DATOS A ESCRIBIR ES:')
            print(filesheet)
            #Script para eliminar los datos dentro de una hoja de calculo de excel
            wb = load_workbook(filesheet)
            print('................................................................')
            print('LAS HOJAS DE CÁLCULO EXISTENTES SON:')
            print(wb.sheetnames)
        
        case 'Trasladar a Excel':
            try:
                dir_actual1 = os.getcwd()
                new_dir_actual1 = dir_actual1.replace('\\','/')
                filesheet1 = f'{new_dir_actual1}/datos.xlsx'
                wb1 = load_workbook(filesheet1)
                sheet = values['hoja'] #extraigo del diccionario, el contenido de la llave hoja.
                
                # Muestra la hoja de cálculo que deseo cambiar
                ws = wb1[sheet]
                print('................................................................')
                print('LA HOJA DE CÁLCULO A LLENAR ES:')
                print(ws)

                #Borro lo que esta dentro de la hoja de excel.
                for row in ws:
                    for cell in row:
                        cell.value = None
                
                wb1.save(filesheet1)

                # AQUI, ES NECESARIO, COLOCAR UN TIEMPO PARA QUE LOS RESULTADOS SURTAN EFECTO
                # Aunque no se ha visto afectado. se sugiere colocar.
                time.sleep(3)

                #Rutina para editar una hoja de excel
                dir_actual2 = os.getcwd()
                new_dir_actual2 = dir_actual2.replace('\\','/')
                filesheet2 = f'{new_dir_actual2}/datos.xlsx'
                wb2 = load_workbook(filesheet2)

                sheet = values['hoja'] #extraigo del diccionario, el contenido de la llave hoja.

                # Muestra la hoja de cálculo que deseo cambiar
                ws1 = wb2[sheet]
                #print(ws1)

                # Pongo dentro del excel lo del diccionario
                for m in dict2:
                    ws1.append(dict2[m])

                wb2.save(filesheet2)
                print('................................................................')
                print('!TRASLADO EXITOSO¡')
            except Exception:
                print('Error: No ingresó el nombre correcto del archivo de word')
                print('o no ingresó el nombre de la hoja de cálculo') 
                print('o no existe hoja de cálculo o archivo de excel de nombre datos.xls.')
                window['indication'].update(value = '!!Debe ingresar nombres correctos y nombre de hoja correctos¡¡')

        case sg.WIN_CLOSED:
            break
window.close()
